
from flask import Flask,render_template,request,jsonify
from chat import get_response
import keras
import numpy as np
import librosa
import pyaudio
import wave
from array import array
import struct
import time
from matplotlib import pyplot as plt
import docx
import IPython.display as ipd
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx2pdf import convert
import math
import time
import datetime
from docx2pdf import convert
import numpy as np
from unittest.main import main
import pymsgbox
from Live_Prediction import *
import eye_detection
import cv2
import dlib
from scipy.spatial import distance as dist
from facial import *
import matplotlib.pyplot as plt
from keras.models import model_from_json
from gaze_tracking import GazeTracking
import stopwatch
from multiprocessing import Process

app = Flask(__name__, static_folder='static')
# to detect the facial region
detector = dlib.get_frontal_face_detector()
predictor = dlib.shape_predictor('shape_predictor_68_face_landmarks.dat')
my_doc1 = docx.Document()
files = []
my_doc = docx.Document()
# Initialize variables
RATE = 24414
CHUNK = 512
RECORD_SECONDS = 20
total_predictions = [] # A list for all predictions in the session.
FORMAT = pyaudio.paInt32
CHANNELS = 1
WAVE_OUTPUT_FILE = "records/output.wav"
#Created the object p of class modelPredictions
emotions = {'0': 'neutral','1': 'calm','2': 'happy','3': 'sad','4': 'angry','5': 'fearful','6': 'disgust','7': 'surprised'}
emo_list = list(emotions.values())
# Open an input channel
p = pyaudio.PyAudio()
stream = p.open(format=FORMAT,
                channels=CHANNELS,
                rate=RATE,
                input=True,
                frames_per_buffer=CHUNK)


background_scripts = {}
@app.get("/")
def indext_get():
    return render_template("index.html")
@app.post("/predict")
def predict ():
    text = request.get_json().get("message")
    #TODO: check if text is valid
    respose = get_response(text)
    message ={"answer":respose}

    return jsonify(message)
class modelPredictions:
    
    def __init__(self, path):
        self.path = path
        #self.file = file

    def load_model(self):
        self.loaded_model = keras.models.load_model(self.path)
        #return self.loaded_model.summary()

    def predictEmotion(self,file):
        data, sr = librosa.load(file)
        mfccs = np.mean(librosa.feature.mfcc(y=data, sr=sr, n_mfcc=40).T, axis=0)
        x = np.expand_dims(mfccs, axis=1)
        x = np.expand_dims(x, axis=0)
        predictions=self.loaded_model.predict(x,use_multiprocessing=True)
        pred_list = list(predictions)
        pred_np = np.squeeze(np.array(pred_list).tolist(), axis=0) # Get rid of 'array' & 'dtype' statments.
        total_predictions.append(pred_np)
        # Present emotion distribution for a sequence (7.1 secs).
        fig = plt.figure(figsize = (7, 2))
        plt.bar(emo_list, pred_np, color = 'darkturquoise')
        plt.ylabel("Probabilty (%)")
        #plt.show()
        fig.savefig('imags/Probabilty (%)_speed.png')
        max_emo = np.argmax(predictions)
        print('max emotion:', emotions.get(max_emo,-1))
        print(100*'-')
        my_doc.add_heading("             The report for Your session", 0)
        my_doc.add_heading("                                                        Probabilties:", 1)
        my_doc.add_heading('                                       From Audio session we detect this Probabilty:',2)    
        my_doc.add_picture('imags/Probabilty (%)_speed.png')     




 #Initialize a non-silent signals array to state "True" in the first 'while' iteration.


def volume(debussy_file):

 ipd.Audio(debussy_file)

# load audio files with librosa
 debussy, sr = librosa.load(debussy_file)
 FRAME_SIZE = 1024
 HOP_LENGTH = 512

 sc_debussy = librosa.feature.spectral_centroid(y=debussy, sr=sr, n_fft=FRAME_SIZE, hop_length=HOP_LENGTH)[0]
 sc_debussy.shape
 frames = range(len(sc_debussy))
 t = librosa.frames_to_time(frames, hop_length=HOP_LENGTH)
 len(t)
 plt.figure(figsize=(7,8))
 # line colour is red
 plt.axhline(y = 4900, color = 'b', linestyle = 'dashed',label = "threshold of normal volume")  
 plt.plot(t, sc_debussy, color='r',label = "Intensity of sound")
 plt.xlabel('Time (sec)')
 plt.ylabel('vloume')
 # plotting the legend
 # Add label
 plt.legend(bbox_to_anchor = (0.5, 1.1), loc = 'upper center')
 plt.savefig('imags/volume of voice.png')
 #plt.show()
 my_doc.add_heading('                                       From Audio session we detect volume summry:',2)    
 my_doc.add_picture('imags/volume of voice.png')     

def is_silent(data):
    # Returns 'True' if below the 'silent' threshold
    return max(data) < 100
@app.route('/session')
def videos_get():
    return render_template("videos (2).html")
@app.route('/facial')
def gaze_blinking_function():
        #variables
     #JAWLINE_POINTS = list(range(0, 17))
     #RIGHT_EYEBROW_POINTS = list(range(17, 22))
     #LEFT_EYEBROW_POINTS = list(range(22, 27))
     #NOSE_POINTS = list(range(27, 36))
     RIGHT_EYE_POINTS = list(range(36, 42))
     LEFT_EYE_POINTS = list(range(42, 48))
     #MOUTH_OUTLINE_POINTS = list(range(48, 61))
     #MOUTH_INNER_POINTS = list(range(61, 68))

     EYE_AR_THRESH = 0.22
     EYE_AR_CONSEC_FRAMES = 2
     #EAR_AVG = 0
     slouch = False
     t_last = time.time()
     count = 0
     COUNTER = 0
     TOTAL = 0
     # counts
     Angry_count = 0
     Disgust_count = 0
     Fear_count = 0
     Happy_count = 0
     Neutral_count = 0
     Sad_count = 0
     Surprise_count = 0
     figure, axis = plt.subplots(2, 2)
     emotion_labels = ['neutral', 'happiness', 'surprise', 'sadness', 'anger', 'disgust', 'fear']
     face_classifier = cv2.CascadeClassifier('haarcascade_frontalface_default.xml')
     # Load model from JSON file
     json_file = open('Saved-Models-facial/model8258.json', 'r')
     loaded_model_json = json_file.read()
     json_file.close()
     classifier = model_from_json(loaded_model_json)

     # Load weights and them to model
     classifier.load_weights('Saved-Models-facial/model8258.h5')

     #cap = cv2.VideoCapture(0)
     neutral_face = []
     neutral_time = []
     happy_time = []
     happy_face = []
     sad_time = []
     sad_face = []
     surprise_time = []
     surprise_face = []
     fear_time = []
     fear_face = []
     disgust_time = []
     disgust_face = []
     angry_time = []
     angry_face = []

     gaze = GazeTracking()
     webcam = cv2.VideoCapture(0)
     dims = get_dims(webcam, res=my_res)
     video_type_cv2 = get_video_type(filename)
     out = cv2.VideoWriter(filename, get_video_type(filename), frames_per_seconds, get_dims(webcam, my_res))
     timer=stopwatch.MyTimer()
     timer.start()


     session_start=time.time()
     time_puse= time.time() - session_start
     while True:
        # We get a new frame from the webcam
        _, frame = webcam.read()
        #_,frame2 = webcam.read()
        # We send this frame to GazeTracking to analyze it
        if int(cv2.__version__.split('.')[0]) >= 3:
            cv_flag = cv2.CASCADE_SCALE_IMAGE
        else:
            cv_flag = cv2.cv.CV_HAAR_SCALE_IMAGE
        gaze.refresh(frame)
        frame = gaze.annotated_frame()
        text = ""
         # convert the frame to grayscale
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        rects = detector(gray, 0)
        #face functions
        datet = str(datetime.datetime.now())
        t = time.localtime()
        current_time = time.strftime("%H:%M:%S", t)

        labels = []
        grayFace = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

        faces = face_classifier.detectMultiScale(grayFace,scaleFactor=1.1,minNeighbors=5,minSize=(30,30),flags=cv_flag)
        #####################Posture#################
        # if int(cv2.__version__.split('.')[0]) >= 3:
        #     cv_flag = cv2.CASCADE_SCALE_IMAGE
        # else:
        #     cv_flag = cv2.cv.CV_HAAR_SCALE_IMAGE
        #
        # faces = faceCascade.detectMultiScale(
        #     gray,
        #     scaleFactor=1.1,
        #     minNeighbors=5,
        #     minSize=(30, 30),
        #     flags=cv_flag
        # )
        ######################################################

        for (x, y, w, h) in faces:
            for rect in rects:
                cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 255), 2)
                roi_gray = grayFace[y:y + h, x:x + w]
                roi_gray = cv2.resize(roi_gray, (48, 48), interpolation=cv2.INTER_AREA)

                cv2.putText(frame, datet, (30, 80), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
                out.write(frame)
                ################Posture####################
                if y > 300 or h > 300:
                    print('slouch')
                    slouch = True
                    #cv2.imwrite("frame%d.jpg" % count, frame)
                    count = count + 1
                else:
                    slouch = False
                    print("No slouching detected")
                ########################################################
                if np.sum([roi_gray]) != 0:
                    roi = roi_gray.astype('float') / 255.0
                    roi = img_to_array(roi)
                    roi = np.expand_dims(roi, axis=0)
                    prediction = classifier.predict(roi)[0]
                    label = emotion_labels[prediction.argmax()]
                    if (prediction.argmax() == 4):
                        Angry_count = Angry_count + 1
                        angry_face.append(Angry_count)
                        angry_time.append(current_time)
                        #plot1 = plt.figure(1)
                        #plt.plot(angry_time, angry_face)
                        #plt.xticks(rotation=45)
                        #plt.title('Angry')
                        #plt.xlabel('Time')
                        #plt.ylabel('count')
                        #plot1.savefig('imags/angry_report.png')



                    elif (prediction.argmax() == 5):
                        Disgust_count = Disgust_count + 1
                        disgust_face.append(Disgust_count)
                        disgust_time.append(current_time)
                        #plot2 = plt.figure(2)
                        #plt.plot(disgust_time, disgust_face)
                        #plt.xticks(rotation=45)
                        #plt.title('Disgust')
                        #plt.xlabel('Time')
                        #plt.ylabel('count')
                        #plot2.savefig('imags/disgust_report.png')

                    elif (prediction.argmax() == 6):
                        Fear_count = Fear_count + 1
                        fear_face.append(Fear_count)
                        fear_time.append(current_time)
                        #plot3 = plt.figure(3)
                        #plt.plot(fear_time, fear_face)
                        #plt.xticks(rotation=45)
                        #plt.title('Fear')
                        #plt.xlabel('Time')
                        #plt.ylabel('count')
                        #plot3.savefig('imags/fear_report.png')

                    elif (prediction.argmax() == 1):
                        Happy_count = Happy_count + 1
                        happy_face.append(Happy_count)
                        happy_time.append(current_time)
                        #plot4 = plt.figure(4)
                        #plt.plot(happy_time, happy_face)
                        #plt.xticks(rotation=45)
                        #plt.title('Happy')
                        #plt.xlabel('Time')
                        #plt.ylabel('count')
                        #plot4.savefig('imags/happy_report.png')



                    elif (prediction.argmax() == 0):
                        Neutral_count = Neutral_count + 1
                        neutral_face.append(Neutral_count)
                        neutral_time.append(current_time)
                        #plot5 = plt.figure(5)
                        #plt.plot(neutral_time, neutral_face)
                        #plt.xticks(rotation=45)
                        #plt.title('Neutral')
                        #plt.xlabel('Time')
                        #plt.ylabel('count')
                        #plot5.savefig('imags/neutral_report.png')


                    elif (prediction.argmax() == 3):
                        Sad_count = Sad_count + 1
                        sad_face.append(Sad_count)
                        sad_time.append(current_time)
                        #plot6 = plt.figure(6)
                        #plt.plot(sad_time, sad_face)
                        #plt.xticks(rotation=45)
                        #plt.title('Sad')
                        #plt.xlabel('Time')
                        #plt.ylabel('count')
                        #plot6.savefig('imags/sad_report.png')

                    elif (prediction.argmax() == 2):
                        Surprise_count = Surprise_count + 1
                        surprise_face.append(Surprise_count)
                        surprise_time.append(current_time)
                        #plot7 = plt.figure(7)
                        #plt.plot(surprise_time, surprise_face)
                        #plt.xticks(rotation=45)
                        #plt.title('Surprise')
                        #plt.xlabel('Time')
                        #plt.ylabel('count')
                        #plot7.savefig('imags/surprise_report.png')

                    label_position = (x, y - 10)
                    cv2.putText(frame, label, label_position, cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)



                else:
                    cv2.putText(frame, 'No Faces', (30, 80), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)


            #for rect in rects:
               # x = rect.left()
               # y = rect.top()
               # x1 = rect.right()
               # y1 = rect.bottom()
                    # get the facial landmarks
                landmarks = np.matrix([[p.x, p.y] for p in predictor(frame, rect).parts()])
                    # get the left eye landmarks
                left_eye = landmarks[LEFT_EYE_POINTS]
                    # get the right eye landmarks
                right_eye = landmarks[RIGHT_EYE_POINTS]
                    # draw contours on the eyes
                #left_eye_hull = cv2.convexHull(left_eye)
                #right_eye_hull = cv2.convexHull(right_eye)
               # cv2.drawContours(frame, [left_eye_hull], -1, (0, 255, 0), 1) # (image, [contour], all_contours, color, thickness)
               # cv2.drawContours(frame, [right_eye_hull], -1, (0, 255, 0), 1)
                    # compute the EAR for the left eye
                ear_left = eye_detection.eye_aspect_ratio(left_eye)
                    # compute the EAR for the right eye
                ear_right = eye_detection.eye_aspect_ratio(right_eye)
                    # compute the average EAR
                ear_avg = (ear_left + ear_right) / 2.0
                    # detect the eye blink
                if ear_avg < EYE_AR_THRESH:
                        COUNTER += 1
                else:
                        if COUNTER >= EYE_AR_CONSEC_FRAMES:
                            TOTAL += 1
                            print("Eye blinked")
                        COUNTER = 0

            if gaze.is_right():
                 text = "Looking right"#pused
                 #print("time pused")
                 #print(timer.pause())
                 print("Eye diverted")

            elif gaze.is_left():
                 text = "Looking left"#pused
                 #print("time pused")
                 #print(timer.pause())
                 print("Eye diverted")


            elif gaze.is_center():
                 text = "Looking center"#running
                 timer.resume()
                 time_puse= time.time() - time_puse
                 print("Eye centered")
            cv2.putText(frame, text, (90, 60), cv2.FONT_HERSHEY_DUPLEX, 1.6, (147, 58, 31), 2)

            #left_pupil = gaze.pupil_left_coords()
            #right_pupil = gaze.pupil_right_coords()
            #cv2.putText(frame, "Left pupil:  " + str(left_pupil), (90, 130), cv2.FONT_HERSHEY_DUPLEX, 0.9, (147, 58, 31), 1)
            #cv2.putText(frame, "Right pupil: " + str(right_pupil), (90, 165), cv2.FONT_HERSHEY_DUPLEX, 0.9, (147, 58, 31), 1)

            cv2.imshow("Demo", frame)
            #cv2.imshow("Face",frame2)

            if slouch == True and time.time() - t_last > 10:
                pymsgbox.alert('Stop Slouching!', 'PostureFix')
                t_last = time.time()
        if cv2.waitKey(1) == 27:
            session_time=time.time()-session_start
            eye_contact= ((time_puse)/(session_time))
            eye_contact=(1-eye_contact)*100

            print("session time= "+ str((session_time))+" sec")
            print("time puse ="+ str((time_puse))+" sec")
            print("eye contact ="+str(math.floor(eye_contact) )+" %" )
            print("Total blinks =  "+str(TOTAL))
            my_doc1.add_heading('                                      From Facial session we detect this summry:',1)
            my_doc1.add_heading("         Session time= "+ str((session_time))+" sec", 1)
            my_doc1.add_heading("         Time puse ="+ str((time_puse))+" sec", 1)
            my_doc1.add_heading("         Eye contact ="+str(math.floor(eye_contact) )+" %", 1)
            my_doc1.add_heading("         Total blinks =  "+str(TOTAL), 1)            
            break

     # defining labels
     activities_label = ['neutral', 'happiness', 'surprise', 'sadness', 'anger', 'disgust', 'fear']
     activities = [1, 2, 3, 4, 5, 6, 7]
     slices = [Neutral_count, Happy_count, Surprise_count, Sad_count, Angry_count, Disgust_count, Fear_count]

     # color for each label
     colors = ['r', 'y', 'g', 'b', 'brown', 'black', 'orange']

     fig = plt.figure(8)
     plt.bar(activities, slices, tick_label=activities_label,width=0.8, color=colors)
     plt.title('The faces count!')
     #fig.show()    
     fig.savefig('imags/count_report.png')
     
     my_doc1.add_picture('imags/count_report.png')    
     my_doc1.save("report/report_from_session1.docx") 
     combine_all_docx()
     webcam.release()
     cv2.destroyAllWindows()
     
@app.route('/speech')
def speech():
    p1 = modelPredictions(path='SER_model.h5')
    p1.load_model()
    data = array('h', np.random.randint(size = 512, low = 0, high = 500))
# SESSION START
    print("** session started")
 
    tic = time.perf_counter()

    while is_silent(data) == False:
       print("* recording...")
       frames = [] 
       data = np.nan # Reset 'data' variable.

       timesteps = int(RATE / CHUNK * RECORD_SECONDS) # => 339

    # Insert frames to 'output.wav'.
       for i in range(0, timesteps):
            data = array('l', stream.read(CHUNK)) 
            frames.append(data)

            wf = wave.open(WAVE_OUTPUT_FILE, 'wb')
            wf.setnchannels(CHANNELS)
            wf.setsampwidth(p.get_sample_size(FORMAT))
            wf.setframerate(RATE)
            wf.writeframes(b''.join(frames))

    
       p1.predictEmotion(file=WAVE_OUTPUT_FILE)
    
    # Define the last 2 seconds sequence.
       last_frames = np.array(struct.unpack(str(96 * CHUNK) + 'B' , np.stack(( frames[-1], frames[-2], frames[-3], frames[-4],
                                                                            frames[-5], frames[-6], frames[-7], frames[-8],
                                                                            frames[-9], frames[-10], frames[-11], frames[-12],
                                                                            frames[-13], frames[-14], frames[-15], frames[-16],
                                                                            frames[-17], frames[-18], frames[-19], frames[-20],
                                                                            frames[-21], frames[-22], frames[-23], frames[-24]),
                                                                            axis =0)) , dtype = 'b')
       if is_silent(last_frames): # If the last 2 seconds are silent, end the session.
           break

# SESSION END        
    toc = time.perf_counter()
    stream.stop_stream()
    stream.close()
    p.terminate()
    wf.close()
    print('** session ended')

# Present emotion distribution for the whole session.
    total_predictions_np =  np.mean(np.array(total_predictions).tolist(), axis=0)
    fig = plt.figure(figsize = (7, 5))
    plt.bar(emo_list, total_predictions_np, color = 'indigo')
    plt.ylabel("Mean probabilty (%)")
    plt.title("Session Summary")
    #plt.show()
    fig.savefig('imags/Session Summary_speech.png')
    print(f"Emotions analyzed for: {(toc - tic):0.4f} seconds")
    my_doc.add_heading('                                       From Audio session we detect this summry:',1)    
    my_doc.add_picture('imags/Session Summary_speech.png')     
    volume(WAVE_OUTPUT_FILE) 
    my_doc.save("report/report_from_session.docx")  
   
def combine_all_docx():
    master = Document_compose("report/report_from_session.docx")
    composer = Composer(master)
   #filename_second_docx is the name of the second docx file
    doc2 = Document_compose("report/report_from_session1.docx")
    #append the doc2 into the master using composer.append function
    composer.append(doc2)
    #Save the combined docx with a name
    composer.save("report/combined.docx")     
    convert("report/combined.docx", "report/combined.pdf")      


if __name__ == '__main__':
    app.run(debug=True)
