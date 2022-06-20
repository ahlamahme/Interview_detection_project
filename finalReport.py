from docx import Document
import aspose.words as aw
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from volume_meter import *
from docx2pdf import convert



def create_qa(self):
        # Table data in a form of list
        print(self.qa);
        data = ((1, 'Geek 1'),(2, 'Geek 2'),(3, 'Geek 3')) 
        # Adding heading in the 1st row of the table
        table = self.add_table(rows=1, cols=2)
        row = table.rows[0].cells
        row[0].text = 'Question'
        row[1].text = 'Answer'
  
        # Adding data from the list to the table
        for id, name in data:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(id)
            row[1].text = name
    
    

def write_results(qs,uid,contact,mins,blinks,slouch):
        b=""
        user_x= Document('report/report_template.docx')
        p = user_x.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()

        #r.add_text(qa)
        #r.add_break(WD_BREAK.PAGE)

        r.add_text("Question 1 : "+qs['0'] +"\n")
        r.add_picture(uid+'/Session Summary_speech0.png')
        r.add_break(WD_BREAK.PAGE)
        
        r.add_text("Question 1 : "+qs['0']+"\n")
        r.add_picture(uid+'/Session Summary_face0.png')
        r.add_break(WD_BREAK.PAGE)
        
        
        r.add_text("Question 2 : "+qs['1']+"\n")
        r.add_picture(uid+'/Session Summary_speech1.png')
        r.add_break(WD_BREAK.PAGE)
        
        r.add_text("Question 2 : "+qs['1']+"\n")
        r.add_picture(uid+'/Session Summary_face1.png')
        r.add_break(WD_BREAK.PAGE)
       

        r.add_text("Question 3 : "+qs['2'])
        r.add_picture(uid+'/Session Summary_speech2.png')
        r.add_break(WD_BREAK.PAGE)
       
        r.add_text("Question 3 : "+qs['2'])
        r.add_picture(uid+'/Session Summary_face2.png')
        r.add_break(WD_BREAK.PAGE)
      

        r.add_text("Question 4 : "+qs['3'])
        r.add_picture(uid+'/Session Summary_speech3.png')
      
        r.add_text("Question 4 : "+qs['3'])
        r.add_picture(uid+'/Session Summary_face3.png')
        r.add_break(WD_BREAK.PAGE)
      

        r.add_text("Question 5 : "+qs['4'])
        r.add_picture(uid+'/Session Summary_speech4.png')
        r.add_break(WD_BREAK.PAGE)
     
        r.add_text("Question 5 : "+qs['4'])
        r.add_picture(uid+'/Session Summary_face4.png')
        r.add_break(WD_BREAK.PAGE)

        vm = volume_meter(uid+'/temp_audio.wav')
        vm.volume(uid);

        r.add_text("Here is your volume loudness through out your interview session:")
        r.add_picture(uid+'/volume of voice.png')
        r.add_break(WD_BREAK.PAGE)

        #res={'session time':300,'blinks':5,'eye contact':70} #for testing purpose

        r.add_text("Now eyes analysis:")
        r.add_break(WD_BREAK.LINE)
        r.add_picture('static/images/eye.jpg')
        r.add_break(WD_BREAK.LINE)

        if contact>60:
            ec= "your maintained good eye contact with interviewer for "+str(contact)+ " %"
        else:
            ec= "your maintained bad eye contact with interviewer for "+str(contact)+ " %"  

        if blinks/mins < 8:
            b = "you blinked "+str(blinks)+ " times in "+ str(mins) + " mins, that is less than an average human bieng, Not blinking normally can indicate discomfort or shock"
        elif blinks/mins > 12:
            b = "you blinked "+str(blinks)+ " times in "+ str(mins) + " mins, that is more than an average human bieng, Not blinking normally can indicate discomfort or shock"

        r.add_text("your interview session lasted for "+str(mins)+ " mins." )
        r.add_break(WD_BREAK.LINE)
        r.add_text(ec)
        r.add_break(WD_BREAK.LINE)
        r.add_text(b)
        r.add_break(WD_BREAK.LINE)

        r.add_break(WD_BREAK.PAGE)
        #count=1 #just for testing purposes

        if slouch > 0:
            r.add_text("You were detected slouching " + str(slouch) +" times, and here is an evidence :p")
            r.add_break(WD_BREAK.LINE)
            r.add_picture(uid+'/frame0.jpg',width=Inches(6.00), height=Inches(6.00) )
        else:
            r.add_text("Your posture showed that you were confident and engaged in the interview, Great Job." +str(slouch))
            r.add_break(WD_BREAK.LINE)

        r.add_text("Finally, Practice makes perfect. You can always redo this interview session and improve your self.\n")
        r.add_break(WD_BREAK.LINE)
        r.add_text("Oh if there's a question that you don't know how to answer, just ask Sam the chatbot he knows great tips for every question.\n")
        r.add_break(WD_BREAK.LINE)
        r.add_break(WD_BREAK.LINE)
        r.add_text("We wish you good luck and good salary :)")

        user_x.save(r'C:\Users\vip\Documents\gp-master\\'+ uid +'\\r'+uid+'.docx')
        convert(r'C:\Users\vip\Documents\gp-master\\'+ uid +'\\r'+uid+'.docx')
        print("Done")
        return True


if __name__== "__main__":
    qs={'0':'introduce','1':'weakness','2':'strength','3':'salary','4':'money work?',};
    write_results(qs,"0",60,5,1,0)
