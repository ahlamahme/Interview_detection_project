#!/usr/bin/env python
# coding: utf-8

import os
# Keyboard module in Python
import keyboard  
import keras
import numpy as np
import librosa
import pyaudio
import wave
from array import array
import struct
import time
from matplotlib import pyplot as plt
import docx
import IPython.display as ipd
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx2pdf import convert
from flask import Flask, render_template

my_doc = docx.Document()
# Initialize variables
RATE = 24414
CHUNK = 512
RECORD_SECONDS = 20
total_predictions = [] # A list for all predictions in the session.
FORMAT = pyaudio.paInt32
CHANNELS = 1
WAVE_OUTPUT_FILE = "records/output.wav"
#Created the object p of class modelPredictions
emotions = {'0': 'neutral','1': 'calm','2': 'happy','3': 'sad','4': 'angry','5': 'fearful','6': 'disgust','7': 'surprised'}
emo_list = list(emotions.values())
# Open an input channel
p = pyaudio.PyAudio()
stream = p.open(format=FORMAT,
                channels=CHANNELS,
                rate=RATE,
                input=True,
                frames_per_buffer=CHUNK)






class modelPredictions:

    def __init__(self, path):
        self.path = path
        #self.file = file

    def load_model(self):
        self.loaded_model = keras.models.load_model(self.path)
        #return self.loaded_model.summary()

    def predictEmotion(self,file):
        data, sr = librosa.load(file)
        mfccs = np.mean(librosa.feature.mfcc(y=data, sr=sr, n_mfcc=40).T, axis=0)
        x = np.expand_dims(mfccs, axis=1)
        x = np.expand_dims(x, axis=0)
        predictions=self.loaded_model.predict(x,use_multiprocessing=True)
        pred_list = list(predictions)
        pred_np = np.squeeze(np.array(pred_list).tolist(), axis=0) # Get rid of 'array' & 'dtype' statments.
        total_predictions.append(pred_np)
        # Present emotion distribution for a sequence (7.1 secs).
        fig = plt.figure(figsize = (7, 2))
        plt.bar(emo_list, pred_np, color = 'darkturquoise')
        plt.ylabel("Probabilty (%)")
        #plt.show()
        fig.savefig('imags/Probabilty (%)_speed.png')
        max_emo = np.argmax(predictions)
        print('max emotion:', emotions.get(max_emo,-1))
        print(100*'-')
        my_doc.add_heading("             The report for Your session", 0)
        my_doc.add_heading("                                                        Probabilties:", 1)
        my_doc.add_heading('                                       From Audio session we detect this Probabilty:',2)    
        my_doc.add_picture('imags/Probabilty (%)_speed.png')     




 #Initialize a non-silent signals array to state "True" in the first 'while' iteration.


def volume(debussy_file):

 ipd.Audio(debussy_file)

# load audio files with librosa
 debussy, sr = librosa.load(debussy_file)
 FRAME_SIZE = 1024
 HOP_LENGTH = 512

 sc_debussy = librosa.feature.spectral_centroid(y=debussy, sr=sr, n_fft=FRAME_SIZE, hop_length=HOP_LENGTH)[0]
 sc_debussy.shape
 frames = range(len(sc_debussy))
 t = librosa.frames_to_time(frames, hop_length=HOP_LENGTH)
 len(t)
 plt.figure(figsize=(7,8))
 # line colour is red
 plt.axhline(y = 4900, color = 'b', linestyle = 'dashed',label = "threshold of normal volume")  
 plt.plot(t, sc_debussy, color='r',label = "Intensity of sound")
 plt.xlabel('Time (sec)')
 plt.ylabel('vloume')
 # plotting the legend
 # Add label
 plt.legend(bbox_to_anchor = (0.5, 1.1), loc = 'upper center')
 plt.savefig('imags/volume of voice.png')
 #plt.show()
 my_doc.add_heading('                                       From Audio session we detect volume summry:',2)    
 my_doc.add_picture('imags/volume of voice.png')     




def is_silent(data):
    # Returns 'True' if below the 'silent' threshold
    
    return max(data) < 100


def speech():
  
    flag=False
    p1 = modelPredictions(path='SER_model.h5')
    p1.load_model()
    data = array('h', np.random.randint(size = 512, low = 0, high = 500))
    tic = time.perf_counter()
# SESSION START
    print("** session started")
 
   

    while is_silent(data) == False:
       print("* recording...")
       frames = [] 
       data = np.nan # Reset 'data' variable.

       timesteps = int(RATE / CHUNK * RECORD_SECONDS) # => 339

    # Insert frames to 'output.wav'.
       #for i in range(0, timesteps):
   
       while 1: 
          
            data = array('l', stream.read(CHUNK)) 
            frames.append(data)
            wf = wave.open(WAVE_OUTPUT_FILE, 'wb')
            wf.setnchannels(CHANNELS)
            wf.setsampwidth(p.get_sample_size(FORMAT))
            wf.setframerate(RATE)
            wf.writeframes(b''.join(frames))
            
            if keyboard.is_pressed ("esc")  :
                flag=True
               # p1.predictEmotion(file=WAVE_OUTPUT_FILE)  
                break
            
              
           
       
       p1.predictEmotion(file=WAVE_OUTPUT_FILE)
    
    # Define the last 2 seconds sequence.
       last_frames = np.array(struct.unpack(str(96 * CHUNK) + 'B' , np.stack(( frames[-1], frames[-2], frames[-3], frames[-4],
                                                                            frames[-5], frames[-6], frames[-7], frames[-8],
                                                                            frames[-9], frames[-10], frames[-11], frames[-12],
                                                                            frames[-13], frames[-14], frames[-15], frames[-16],
                                                                            frames[-17], frames[-18], frames[-19], frames[-20],
                                                                            frames[-21], frames[-22], frames[-23], frames[-24]),
                                                                            axis =0)) , dtype = 'b')
        # If the last 2 seconds are silent, end the session.
      # if is_silent(last_frames):
       #      pass
      # 
    
# SESSION END 
       if flag==True :            
        toc = time.perf_counter()
        stream.stop_stream()
        stream.close()
        p.terminate()
        wf.close()
        print('** session ended')
        break

# Present emotion distribution for the whole session.
    total_predictions_np =  np.mean(np.array(total_predictions).tolist(), axis=0)
    fig = plt.figure(figsize = (7, 5))
    plt.bar(emo_list, total_predictions_np, color = 'indigo')
    plt.ylabel("Mean probabilty (%)")
    plt.title("Session Summary")
     #plt.show()
    fig.savefig('imags/Session Summary_speech.png')
    print(f"Emotions analyzed for: {(toc - tic):0.4f} seconds")
    my_doc.add_heading('                                       From Audio session we detect this summry:',1)    
    my_doc.add_picture('imags/Session Summary_speech.png')     
    volume(WAVE_OUTPUT_FILE) 
    my_doc.save("report/report_from_session.docx")  
    # convert("report/report_from_session.docx", "report/report_from_session.pdf")  
    
      
#def get_report(doc):
 #   doc.save("report/report_from_session.docx")
def combine_all_docx():
    master = Document_compose("report/report_from_session.docx")
    composer = Composer(master)
   #filename_second_docx is the name of the second docx file
    doc2 = Document_compose("report/report_from_session1.docx")
    #append the doc2 into the master using composer.append function
    composer.append(doc2)
    #Save the combined docx with a name
    composer.save("report/combined.docx")     
    convert("report/combined.docx", "report/combined.pdf")  
#speech()    